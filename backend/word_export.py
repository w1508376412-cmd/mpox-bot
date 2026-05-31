"""Markdown转Word文档模块"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
from datetime import datetime
from reference_sources import select_reference_sources


def markdown_to_docx(
    markdown_text: str,
    question: str = "",
    user_name: str = "",
    antiviral_id: str = "",
    sources: list = None
) -> BytesIO:
    """
    将markdown文本转换为Word文档

    Args:
        markdown_text: markdown格式的文本
        question: 用户问题
        user_name: 用户姓名
        antiviral_id: 抗病毒编号
        sources: 参考来源列表

    Returns:
        BytesIO对象（Word文档）
    """
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    # 标题
    title = doc.add_heading('猴痘知识问答记录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 用户信息
    if user_name or antiviral_id:
        info_para = doc.add_paragraph()
        info_para.add_run(f'用户姓名：{user_name}    ').bold = False
        info_para.add_run(f'抗病毒编号：{antiviral_id}').bold = False
        info_para.add_run(f'\n咨询时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    doc.add_paragraph('─' * 40)

    # 用户问题
    if question:
        q_heading = doc.add_heading('问题', level=1)
        doc.add_paragraph(question)

    # 回答标题
    doc.add_heading('回答', level=1)

    # 解析markdown并添加到文档
    parse_markdown_to_docx(doc, markdown_text)

    # 参考来源
    if sources and len(sources) > 0:
        doc.add_paragraph('')
        doc.add_heading('参考来源', level=1)
        for src in select_reference_sources(sources):
            source_name = src.get('source', '') if isinstance(src, dict) else src.source
            publish_date = src.get('publish_date', '') if isinstance(src, dict) else src.publish_date
            url = src.get('url', '') if isinstance(src, dict) else src.url
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'{source_name} ({publish_date})')
            if url:
                p.add_run(f'\n  {url}').font.size = Pt(9)

    # 免责声明
    doc.add_paragraph('')
    doc.add_paragraph('─' * 40)
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        '医疗免责声明：本内容仅提供猴痘/mpox健康科普信息，不能替代医生诊断或治疗。'
        '如出现新发或原因不明皮疹、发热、淋巴结肿大，或有可疑接触史，'
        '请咨询医疗机构或当地疾控部门。'
    )
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 保存到BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def parse_markdown_to_docx(doc, markdown_text: str):
    """
    简单的markdown解析器，将markdown转换为Word段落

    支持：
    - # ## ### 标题
    - **加粗**
    - - 或 * 列表
    - 1. 数字列表
    - 普通段落
    """
    lines = markdown_text.strip().split('\n')

    for line in lines:
        line = line.rstrip()

        if not line.strip():
            continue

        # 一级标题
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        # 二级标题
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        # 三级标题
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        # 无序列表
        elif re.match(r'^\s*[-*]\s+', line):
            content = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
        # 有序列表
        elif re.match(r'^\s*\d+\.\s+', line):
            content = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, content)
        # 普通段落
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)


def add_formatted_text(paragraph, text: str):
    """
    处理段落中的格式化文本（如加粗）

    Args:
        paragraph: docx段落对象
        text: 包含markdown格式的文本
    """
    # 匹配**加粗**文本
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # 普通文本
            paragraph.add_run(part)
